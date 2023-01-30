from __future__ import print_function
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload
import csv
import email
import json
import sys
from typing import Tuple
import pandas as pd
import openai

import docx
from docx import Document
from docx.shared import Inches
import google.auth


# import pymongo as pymongo


class Scribe:
    table = ["Primary Owner", "APN/PIN", "Name", "Primary Owner Last Name", "Primary Owner First Name", "Site Address",
             "City", "State", "Zip", "Acreage (Deeded)", "Acreage (Calc)", "Name", "Name",
             "Total Acres - Same Land Owner", "Owner", "City", "County", "State", "Zip Code", "Company Name",
             "Street Address", "City", "State", "Zip", "Muni Id", "Muni Name", "Buildable Area (Acres)", "Notes"]

    def __init__(self, csv_path: str, client_params: str):
        self.csv_path = csv_path
        self.client_params = client_params
        openai.organization = "org-qcHPiKIimtg6ssjx0Xla5AGH"
        openai.api_key = "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"

    def getPropertyData(self):
        # Open the file and read the data into a list
        with open(self.csv_path) as file:
            reader = csv.reader(file)
            data = [row for row in reader]

        # Get the headers and indexes of the columns we want
        headers = data[0]
        name = headers.index("Name")
        # print(name)
        company_index = headers.index('Owner')
        owner_first_name_index = headers.index('Primary Owner First Name')
        owner_last_name_index = headers.index('Primary Owner Last Name')
        site_address_index = headers.index('Site Address')
        owner_address_index = headers.index('Street Address')
        acreage_deeded_index = headers.index('Acreage (Deeded)')
        acreage_calc_index = headers.index('Acreage (Calc)')
        total_acreage_index = headers.index('Total Acres - Same Land Owner')
        city_index = headers.index('City')
        zip_index = headers.index('Zip')
        state_index = headers.index('State')

        # Loop through the data and extract the information
        for row in data[1:]:
            company = row[company_index]
            owner_name = row[owner_first_name_index] + " " + row[owner_last_name_index]
            site_address = row[site_address_index]
            owner_address = row[owner_address_index]
            city = row[city_index]
            state = row[state_index]
            zip = row[zip_index]

            # Print the information
            if owner_name == "":
                to = company
            elif company == "":
                to = owner_name
            else:
                to = name
            if row[acreage_calc_index] == "":
                acreage = row[acreage_deeded_index]
            elif acreage_deeded_index == "":
                acreage = row[acreage_calc_index]
            else:
                acreage = row[total_acreage_index]
            # sys.stdout.write(acreage + " + ")
            self.generate_letter(
                f"My Name: Warren Kritko, Owner Name: {to}, Site Address: {site_address}, City: {city}, State: {state}, Zip: {zip}, Acreage: {acreage}, Owner Address: {owner_address}, Full property information: {row}")

    def generate_letter(self, param_plain_text: str) -> dict:
        response = openai.Completion.create(
            model="text-davinci-003",
            prompt=self.client_params + param_plain_text,
            temperature=0,
            max_tokens=1000
        )
        print(f"{response['choices'][0]['text']}\n")
        return response

    def save_response(self, response: dict):
        response_str = json.dumps(response)
        with open("response.json", "w") as json_file:
            json.dump(response_str, json_file)

    def extract_email_info(self, msg: email.message.EmailMessage) -> Tuple[str, str]:
        name = msg['From']
        carat_email = msg['Return-Path']
        decarated_email = carat_email[1:-1]
        return name, decarated_email

    def create_document(self):
        document = Document()

        document.add_heading('Document Title', 0)
        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')

        p.add_run('italic.').italic = True
        document.add_heading('Heading, level 1', level=1)

        document.add_paragraph('Intense quote', style='Intense Quote')
        document.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )

        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        # document.add_picture('monty-truth.png', width=Inches(1.25))

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'

        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        document.add_page_break()

        document.save('demo.docx')

    def upload_basic(self, file: str):
        """Insert new file.
        Returns : Id's of the file uploaded

        Load pre-authorized user credentials from the environment.
        TODO(developer) - See https://developers.google.com/identity
        for guides on implementing OAuth2 for the application.
        """
        creds, _ = google.auth.default()

        try:
            # create drive api client
            service = build('drive', 'v3', credentials=creds)

            file_metadata = {'name': file}
            media = MediaFileUpload(file,
                                    mimetype='image/jpeg')
            # pylint: disable=maybe-no-member
            file = service.files().create(body=file_metadata, media_body=media,
                                          fields='id').execute()
            print(F'File ID: {file.get("id")}')

        except HttpError as error:
            print(F'An error occurred: {error}')
            file = None

        return file.get('id')


import subprocess

lpr = subprocess.Popen("/usr/bin/lpr", stdin=subprocess.PIPE)

email_client = Scribe(csv_path="/Users/collindrake/Downloads/export.csv",
                      client_params="Write me a letter to inquire with a landowner about buying or leasing their "
                                    "property for a potential solar farm.",
                      )
# email_client.generate_letter(email_client.getPropertyData())

email_client.getPropertyData()

upload_basic(email_client.create_document())

"""
while True:
    try:
    
        client = pymongo.MongoClient(
            "mongodb+srv://frac:frac1@cluster0.ithicjt.mongodb.net/?retryWrites=true&w=majority")
        db = client.test
        print(f"DB: {db}")

        outer_name = ""

        outer_response = email_client.generate_letter(outer_plain_text)
        print(f"response: {outer_response}")

        email_client.save_response(outer_response)

        outer_json_text = json.loads(json.dumps(outer_response))
        print(f"json_text: {outer_json_text}")

        outer_text = outer_json_text['choices'][0]['text']
        print(f"text: {outer_text}")

        # Create the email message
        outer_message = f"Dear {outer_name},\n\n{outer_text}\n"

    except IndexError as e:
        print(f"There are no emails to process currently, checking again in one minute.")
        time.sleep(60)
    except Exception as j:
        print(f"{j}")
        time.sleep(60)
"""
